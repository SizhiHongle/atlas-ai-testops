"""Migrate Atlas DOCX design documents to the accepted Python backend baseline."""

from __future__ import annotations

import re
from datetime import UTC, datetime
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentType
from docx.text.paragraph import Paragraph

REPOSITORY_ROOT = Path(__file__).resolve().parents[2]
DOCUMENTS_DIR = REPOSITORY_ROOT / "documents"

VERSION_METADATA: dict[str, tuple[str, str]] = {
    "AI_Test_Platform_Implementation_Plan_v1.1.docx": (
        "AI 测试平台落地方案 v1.1",
        "Accepted baseline",
    ),
    "Atlas_AI用例与浏览器Agent工作流功能设计及实现对齐稿_v0.3.docx": (
        "Atlas AI 用例、编排画布与浏览器 Agent 工作流功能设计及实现对齐稿 v0.3",
        "Aligned baseline",
    ),
    "Atlas_任务中心与批量执行控制面功能设计及实现对齐稿_v0.2.docx": (
        "Atlas 任务中心与批量执行控制面功能设计及实现对齐稿 v0.2",
        "Draft - architecture baseline accepted",
    ),
    "Atlas_数据预加载与原子组件编排功能设计对齐稿_v0.2.docx": (
        "Atlas 数据预加载与原子组件编排功能设计对齐稿 v0.2",
        "Draft - architecture baseline accepted",
    ),
    "Atlas_洞察中心功能设计及实现对齐稿_v0.2.docx": (
        "Atlas 洞察中心功能设计及实现对齐稿 v0.2",
        "Draft - architecture baseline accepted",
    ),
    "Atlas_现场与浏览器实时执行功能设计及实现对齐稿_v0.2.docx": (
        "Atlas 现场与浏览器实时执行功能设计及实现对齐稿 v0.2",
        "Draft - architecture baseline accepted",
    ),
    "Atlas_结果中心功能设计及实现对齐稿_v0.2.docx": (
        "Atlas 结果中心功能设计及实现对齐稿 v0.2",
        "Draft - architecture baseline accepted",
    ),
    "Atlas_身份与测试账号体系技术设计_v1.1.docx": (
        "Atlas 身份与测试账号体系技术设计 v1.1",
        "Accepted baseline",
    ),
}

COVER_DATE_DOCUMENTS = {
    "Atlas_AI用例与浏览器Agent工作流功能设计及实现对齐稿_v0.3.docx",
    "Atlas_任务中心与批量执行控制面功能设计及实现对齐稿_v0.2.docx",
    "Atlas_数据预加载与原子组件编排功能设计对齐稿_v0.2.docx",
}

CODE_REPLACEMENTS: dict[str, dict[int, str]] = {
    "Atlas_AI用例与浏览器Agent工作流功能设计及实现对齐稿_v0.3.docx": {
        119: """class WorkflowDraft(BaseModel):
    schema_version: Literal[\"atlas.workflow-draft/0.1\"]
    id: str
    test_case_id: str
    semantic_revision: int
    layout_revision: int
    graph: WorkflowGraph
    layout: dict[str, NodeLayout]
    intent_version_ref: ExactVersionRef
    updated_by: Literal[\"ai\", \"human\"]

class WorkflowEdge(BaseModel):
    id: str
    source_node_id: str
    source_port: str
    target_node_id: str
    target_port: str
    semantic_type: str
    kind: Literal[\"data\", \"control\"]
    mapping: Literal[\"direct\"] = \"direct\"""",
        124: """class WorkflowPatch(BaseModel):
    patch_id: str
    client_mutation_id: str
    base_semantic_revision: int
    source: Literal[\"ai\", \"human\"]
    operations: tuple[DraftOperation, ...]
    rationale_summary: str | None = None""",
        155: """class TestIRVersion(BaseModel):
    schema_version: Literal[\"atlas.test-ir/0.2\"]
    case_version_id: str
    test_intent_version_id: str
    requirement_refs: tuple[SourceRequirementRef, ...]
    actors: tuple[ActorContract, ...]
    fixture: FixtureContract
    workflow: WorkflowGraphIR
    surfaces: tuple[SurfaceRef, ...]
    variables: dict[str, ValueSource]
    assertions: tuple[AssertionSpec, ...]
    evidence_policy: EvidencePolicy
    recovery_policy: RecoveryPolicy
    outcome_policy: OutcomePolicy
    required_features: tuple[str, ...]
    content_digest: str""",
        162: """class ObjectiveBudget(BaseModel):
    timeout_ms: int
    max_agent_turns: int
    max_browser_actions: int
    max_recoveries: int

class ObjectiveIR(BaseModel):
    id: str
    actor_slot: str
    goal: str
    surface_scope: tuple[str, ...]
    allowed_agent_actions: tuple[BrowserCapability, ...]
    required_transitions: tuple[ActionIntent, ...]
    entry_assertions: tuple[str, ...]
    completion_assertions: tuple[str, ...]
    budget: ObjectiveBudget
    side_effect: Literal[\"read_only\", \"test_data_mutation\"]""",
        169: """class SurfaceContractVersion(BaseModel):
    surface_key: Literal[\"crm.customer.list\"]
    version: Literal[\"1.3.0\"]
    routes: dict[str, RouteContract]
    elements: dict[str, TargetContract]
    content_digest: str""",
        204: """class TargetCandidate(BaseModel):
    target_ref: str
    element_key: str | None = None
    role: str | None = None
    accessible_name: str | None = None
    confidence: float
    semantic_fingerprint: str

class BrowserObservation(BaseModel):
    observation_ref: str
    page_revision: str
    route_key: str | None = None
    title: str
    target_candidates: tuple[TargetCandidate, ...]
    untrusted_page_summary: str
    next_step_nonce: str""",
        262: """class AgentGateway(Protocol):
    async def next_action(self, request: AgentTurnRequest) -> AgentDecision: ...

AgentDecision = ToolDecision | FinishDecision""",
    },
    "Atlas_任务中心与批量执行控制面功能设计及实现对齐稿_v0.2.docx": {
        116: """class TaskPlanVersion(BaseModel):
    schema_version: Literal[\"atlas.task-plan/0.1\"]
    task_key: str
    version: str
    project_id: str
    iteration_id: str | None = None
    case_selector: CaseSelector
    matrix: MatrixDefinition
    execution_policy: ExecutionPolicy
    triggers: tuple[TriggerDefinition, ...]
    gate_policy: GatePolicy
    notification_policy: NotificationPolicy
    status: Literal[\"draft\", \"published\", \"deprecated\"]
    content_digest: str""",
        122: """class PinnedCaseSelector(BaseModel):
    kind: Literal[\"pinned\"]
    case_version_ids: tuple[str, ...]

class QueryAtRunCaseSelector(BaseModel):
    kind: Literal[\"query_at_run\"]
    project_id: str
    iteration_id: str | None = None
    tags: tuple[str, ...] = ()
    risk: tuple[Literal[\"high\", \"critical\"], ...] = ()
    status: Literal[\"published\"] = \"published\"

CaseSelector = PinnedCaseSelector | QueryAtRunCaseSelector""",
        178: """class ExecutionUnitManifest(BaseModel):
    unit_key: str
    case_version_id: str
    execution_profile_version_id: str
    fixture_blueprint_version_id: str
    identity_profile_version_id: str
    environment_id: str
    browser_profile_version_id: str
    data_profile_version_id: str

class TaskRunManifest(BaseModel):
    task_run_id: str
    task_plan_version_id: str
    trigger_fingerprint: str
    project_id: str
    iteration_id: str | None = None
    units: tuple[ExecutionUnitManifest, ...]
    policy_digests: dict[str, str]
    compiler_version: str
    manifest_hash: str""",
    },
    "Atlas_数据预加载与原子组件编排功能设计对齐稿_v0.2.docx": {
        103: """class Effect(StrEnum):
    READ = \"read\"
    CREATE = \"create\"
    UPDATE = \"update\"
    DELETE = \"delete\"
    WAIT = \"wait\"

class HandlerRef(BaseModel):
    module: str
    callable: str
    artifact_digest: str

class DataAtomVersion(BaseModel):
    atom_key: str
    version: str
    input_schema: JsonSchema202012
    output_schema: JsonSchema202012
    ports: tuple[PortDefinition, ...]
    effect: Effect
    required_capabilities: tuple[str, ...]
    retry_policy: RetryPolicy
    idempotency_policy: IdempotencyPolicy
    postconditions: tuple[Postcondition, ...]
    resource_policy: ResourcePolicy | None = None
    cleanup_policy: CleanupContract | None = None
    reconcile_policy: ReconcileContract | None = None
    handler: HandlerRef""",
        105: """class FixtureContext(Protocol):
    connector: ApprovedConnectorClient
    resources: ResourceLedgerClient
    identity: ActorLeaseView
    artifacts: EncryptedArtifactClient
    idempotency_key: str
    execution: ReadonlyExecutionContext""",
        115: """class ResourceRef(BaseModel, Generic[ResourceType]):
    resource_type: ResourceType
    ledger_id: str
    environment_id: str
    connector_id: str
    ownership: Literal[\"created\", \"adopted\", \"leased\"]
    opaque_ref: str""",
    },
    "Atlas_现场与浏览器实时执行功能设计及实现对齐稿_v0.2.docx": {
        149: """class ActionProposal(BaseModel):
    proposal_id: str
    unit_attempt_id: str
    step_ref: str
    target: ActionTarget
    action: Literal[\"click\", \"fill\", \"select\", \"navigate\", \"upload\", \"download\"]
    input: ActionInput | None = None
    expected_observation: tuple[ObservationPredicate, ...]
    evidence_plan: EvidenceCapturePlan
    risk: Literal[\"read\", \"input\", \"navigate\", \"file\", \"destructive\"]
    basis_refs: tuple[str, ...]""",
        165: """class ActionGrant(BaseModel):
    grant_id: str
    proposal_hash: str
    unit_attempt_id: str
    browser_session_id: str
    page_id: str
    control_epoch: int
    fencing_token: int
    allowed_adapter: str
    expires_at: datetime
    max_executions: Literal[1] = 1
    policy_digest: str""",
        259: """Controller = AgentController | HumanController

class LiveSnapshot(BaseModel):
    unit_attempt: UnitAttemptProjection
    browser: BrowserProjection
    control: ControlProjection
    active_action: ActionProjection | None = None
    pages: tuple[PageProjection, ...]
    oracle: OracleProjection
    evidence: EvidenceProjection
    cursor: str

class ExecuteGrantedAction(BaseModel):
    grant: ActionGrant
    proposal: ActionProposal
    expected_browser_revision: str""",
    },
    "Atlas_身份与测试账号体系技术设计_v1.1.docx": {
        93: """class AccountLifecycle(StrEnum):
    DRAFT = \"draft\"
    PROVISIONING = \"provisioning\"
    ACTIVE = \"active\"
    SUSPENDED = \"suspended\"
    RETIRING = \"retiring\"
    RETIRED = \"retired\"

class AccountHealth(StrEnum):
    UNKNOWN = \"unknown\"
    HEALTHY = \"healthy\"
    DEGRADED = \"degraded\"
    QUARANTINED = \"quarantined\"""",
        158: """class ProviderCapability(StrEnum):
    ACCOUNT_DISCOVER = \"account.discover\"
    ACCOUNT_READ = \"account.read\"
    ACCOUNT_PROVISION = \"account.provision\"
    AUTH_PASSWORD = \"auth.password\"
    AUTH_OAUTH2 = \"auth.oauth2\"
    AUTH_OIDC = \"auth.oidc\"
    AUTH_SAML_SSO = \"auth.saml_sso\"
    AUTH_MFA_TOTP = \"auth.mfa.totp\"
    AUTH_MANUAL_BOOTSTRAP = \"auth.manual_bootstrap\"

class CapabilityDescriptor(BaseModel):
    name: ProviderCapability
    version: str
    mode: Literal[\"native_api\", \"browser\", \"webhook\", \"polling\", \"manual\"]
    constraints: CapabilityConstraints | None = None""",
        161: """class IdentityProviderAdapter(Protocol):
    def manifest(self) -> AdapterManifest: ...
    async def probe(self, context: AdapterContext) -> ProbeResult: ...
    async def negotiate(
        self,
        context: AdapterContext,
        requirement: CapabilityRequirement,
    ) -> NegotiatedCapabilities: ...
    async def health(self, context: AdapterContext) -> ProviderHealth: ...""",
        185: """class ExecutionIdentityGrant(BaseModel):
    tenant_id: str
    project_id: str
    environment_id: str
    task_run_id: str
    execution_unit_id: str
    allowed_tool_names: tuple[str, ...]
    allowed_provider_connections: tuple[str, ...]
    allowed_account_lease_id: str | None = None
    allowed_origins: tuple[str, ...]
    expires_at: datetime
    nonce: str""",
        192: """class LoginSessionReady(BaseModel):
    status: Literal[\"ready\"]
    browser_context_ref: str
    expires_at: datetime

class LoginSessionManualAction(BaseModel):
    status: Literal[\"manual_action_required\"]
    action_ticket_id: str
    expires_at: datetime
    safe_reason: str

EnsureLoginSessionResult = LoginSessionReady | LoginSessionManualAction""",
        331: '''# pyproject.toml
[tool.mypy]
python_version = "3.14"
strict = true
warn_unreachable = true
no_implicit_reexport = true

[tool.ruff]
target-version = "py314"

[tool.pytest.ini_options]
asyncio_mode = "auto"''',
        333: """class AccountSource(StrEnum):
    ATLAS_MANAGED = \"atlas_managed\"
    EXTERNAL_SYNCED = \"external_synced\"
    EXTERNAL_DELEGATED = \"external_delegated\"

class AccountLease(BaseModel):
    id: str
    execution_unit_id: str
    account_id: str
    slot_id: str
    status: Literal[\"active\", \"released\", \"expired\", \"revoked\"]
    fencing_token: int
    acquired_at: datetime
    expires_at: datetime
    last_heartbeat_at: datetime
    released_at: datetime | None = None
    release_reason: str | None = None""",
        335: """class IssueSecretGrantInput(BaseModel):
    lease_id: str
    fencing_token: int
    purpose: Literal[\"login\", \"refresh_session\", \"rotate_credential\"]
    worker_identity: str
    allowed_origins: tuple[str, ...]

class SecretGrant(BaseModel):
    grant_ref: str
    expires_at: datetime
    max_redemptions: Literal[1] = 1""",
        337: """class AdapterError(BaseModel):
    code: Literal[
        \"configuration_invalid\",
        \"capability_unsupported\",
        \"authentication_failed\",
        \"credential_expired\",
        \"manual_action_required\",
        \"account_locked\",
        \"rate_limited\",
        \"provider_unavailable\",
        \"network_timeout\",
        \"internal_error\",
    ]
    category: str
    operation: str
    safe_message: str
    retryable: bool
    retry_after_ms: int | None = None
    provider_code: str | None = None
    request_id: str""",
    },
}

TEXT_REPLACEMENTS: tuple[tuple[str, str], ...] = (
    ("IMPLEMENTATION BLUEPRINT  /  V1.0", "IMPLEMENTATION BLUEPRINT  /  V1.1"),
    ("ATLAS  /  TECHNICAL DESIGN  /  V1.0", "ATLAS  /  TECHNICAL DESIGN  /  V1.1"),
    ("QUALITY INTELLIGENCE  /  V0.1", "QUALITY INTELLIGENCE  /  V0.2"),
    ("RESULT INTELLIGENCE  /  V0.1", "RESULT INTELLIGENCE  /  V0.2"),
    ("功能设计及实现对齐稿  v0.1", "功能设计及实现对齐稿  v0.2"),
    ("功能设计对齐稿  v0.1", "功能设计对齐稿  v0.2"),
    ("功能设计及实现对齐稿  /  v0.1", "功能设计及实现对齐稿  /  v0.2"),
    ("文档状态：方案评审稿", "文档状态：架构基线已接受"),
    ("状态：Draft / 架构基线", "状态：Accepted / 架构基线"),
    ("版本：v1.0", "版本：v1.1"),
    ("版本：v0.1", "版本：v0.2"),
    ("日期：2026-07-11", "日期：2026-07-13"),
    ("日期：2026-07-12", "日期：2026-07-13"),
    (
        "2026-07-12  ·  Architecture / Product / Security",
        "2026-07-13  ·  Architecture / Product / Security",
    ),
    ("技术栈：Python + Python 3.14", "技术栈：Python 3.14 + FastAPI + Pydantic"),
    (
        "run_group / test_run / run_node / step_attempt",
        "task_plan_version / task_run / execution_unit / unit_attempt",
    ),
    ("task_run_item", "execution_unit"),
    ("case_run / case_attempt", "execution_unit / unit_attempt"),
    (
        "debug_run / task_run_snapshot / execution_unit / execution_contract / execution_unit / unit_attempt / browser_session",
        "debug_run / task_run_snapshot / execution_unit / unit_attempt / execution_contract / browser_session",
    ),
    ("caseRunId", "executionUnitId"),
    (
        "POST /v1/case-runs                 // Idempotency-Key required",
        "GET  /v1/task-runs/{id}/units",
    ),
    ("GET  /v1/case-runs/{id}/events", "GET  /v1/execution-units/{id}/events"),
    ("GET  /v1/case-runs/{id}/evidence", "GET  /v1/execution-units/{id}/evidence"),
    ("atlas.case.attempt.completed", "atlas.case.unit_attempt.completed"),
    ("atlas.case.run.completed", "atlas.case.execution_unit.completed"),
    ("atlas.case.run.requested", "atlas.case.execution_unit.requested"),
    ("atlas.task.unit.attempt_completed", "atlas.task.unit_attempt.completed"),
    ("first_attempt_pass", "first_unit_attempt_pass"),
    ("重试 attempt 不进入业务幂等键", "Activity retry 序号不进入业务幂等键"),
    (
        "同一 UnitAttempt 内的新 Activity attempt",
        "同一 UnitAttempt 内的新 Activity retry",
    ),
    ("最大 Attempts", "最大 retry attempts"),
    ("保留全部 Attempts", "保留全部 retry attempts"),
    ("UNIT 与 ATTEMPT 分开", "EXECUTION UNIT 与 UNITATTEMPT 分开"),
    ("06 / ATTEMPT SEAL", "06 / UNITATTEMPT SEAL"),
    ("22 / ATTEMPT & EVIDENCE", "22 / UNITATTEMPT & EVIDENCE"),
    ("/results/attempts/:unitAttemptId", "/results/unit-attempts/:unitAttemptId"),
    ("/attempts/", "/unit-attempts/"),
    ("tenant/attempt/session", "tenant/unit_attempt/session"),
    ("tenant + attempt scope", "tenant + unit_attempt scope"),
    ("tenant + attempt 唯一范围", "tenant + unit_attempt 唯一范围"),
    ("atlas.live.attempt.", "atlas.live.unit_attempt."),
    ('"subject": "attempts/', '"subject": "unit-attempts/'),
    ("按 attempt ordinal", "按 UnitAttempt ordinal"),
    ("tenant/project/task/unit/attempt", "tenant/project/task/unit/unitAttempt"),
    ("unit_unit_attempt", "unit_attempt"),
    ("unit-unit-attempt", "unit-attempt"),
    ("UNIT_UNIT_ATTEMPT", "UNIT_ATTEMPT"),
    (
        "RunGroup / Run / NodeAttempt / Event / Evidence / Failure",
        "TaskPlanVersion / TaskRun / ExecutionUnit / UnitAttempt / Event / Evidence / Failure",
    ),
    ("NodeRun / UnitAttempt", "DataNodeRun / DataNodeAttempt"),
    ("UnitAttempt 不改变逻辑幂等键", "DataNodeAttempt 不改变逻辑幂等键"),
    (
        "DataSetupRun       1 ── N NodeRun ── N NodeAttempt",
        "DataSetupRun       1 ── N DataNodeRun ── N DataNodeAttempt",
    ),
    (
        "NodeRun            1 ── N ResourceRecord",
        "DataNodeRun        1 ── N ResourceRecord",
    ),
    ("Fastify + JSON Schema", "FastAPI + Pydantic + JSON Schema"),
    ("NestJS + Fastify，或纯 Fastify", "FastAPI + Pydantic"),
    ("Zod / TypeBox", "Pydantic"),
    ("TypeBox/Zod", "Pydantic"),
    ("pg 显式事务", "Psycopg 3 显式事务"),
    (
        "建议 Python 3.14 + Mypy strict；启用 noUncheckedIndexedAccess。",
        "建议 Python 3.14 + Mypy strict；启用 warn_unreachable 与 strict_optional。",
    ),
    ("Schema 与生成 TS 类型", "Pydantic 模型与导出的 JSON Schema"),
    (
        "所有外部输入使用 schema validation（如 Pydantic）；领域层不接受 unknown 直接透传。",
        "所有外部输入使用 Pydantic 校验；领域层不接受未经校验的 dict 或 Any 直接透传。",
    ),
    (
        "按团队习惯；OpenAPI 与运行时 Schema 同源",
        "FastAPI 原生生成 OpenAPI；Pydantic 与运行时校验同源",
    ),
    ("BrowserAttempt Activity", "BrowserExecution Activity"),
    (
        "Activity 失败后的自动 UnitAttempt 会从初始状态重新执行",
        "Activity 失败后的自动重试会从初始状态重新执行",
    ),
    (
        "// 这不是“完全成功”；严格 Gate 必须拒绝",
        "// Not fully successful; a strict Gate must reject it.",
    ),
    ("// 清理依赖", "// cleanup dependency"),
    ("// 仅诊断", "// diagnostic only"),
    ("# 原始现场与证据永久保留", "# retain the original live state and evidence"),
    ("# 当前现场", "# current live state"),
    ("# 管理与运行 API", "# management and runtime API"),
    ("# 安全工具门面", "# safe tool facade"),
    ("# 回收、健康、漂移", "# reclamation, health, drift"),
    ("# 确定性认证", "# deterministic authentication"),
    ("# SaaS 隔离执行", "# isolated SaaS execution"),
    ("# 纯领域模型与策略", "# pure domain models and policies"),
    ("# Adapter 接口", "# Adapter protocol"),
    ("# 统一合约套件", "# shared contract suite"),
    ("# secret grant、redaction、token", "# secret grants, redaction, tokens"),
    ("TypeScript / Node.js", "Python 3.14"),
    ("Node.js / TypeScript", "Python 3.14"),
    ("Node.js · TypeScript", "Python 3.14"),
    ("TypeScript + Node.js + PostgreSQL", "Python 3.14 + FastAPI + PostgreSQL"),
    (
        "TypeScript + LLM Gateway + Postgres",
        "Python + FastAPI + LLM Gateway + PostgreSQL",
    ),
    ("NestJS + Fastify / TypeScript", "FastAPI / Python 3.14"),
    ("Fastify/NestJS + TypeScript", "FastAPI / Python 3.14"),
    ("NestJS 或 Fastify + TypeScript", "FastAPI + Python 3.14"),
    ("TypeScript + JSON Schema 2020-12", "Pydantic + JSON Schema 2020-12"),
    ("TypeScript + JSON Schema", "Pydantic + JSON Schema"),
    ("Temporal TypeScript SDK", "Temporal Python SDK"),
    ("Temporal TypeScript Cancellation", "Temporal Python Cancellation"),
    ("Node.js 隔离进程/容器", "Python 隔离进程/容器"),
    ("签名 Node.js 包", "签名 Python wheel 或 OCI Artifact"),
    ("Node.js 24 LTS + TypeScript strict", "Python 3.14 + Mypy strict"),
    ("TypeScript strict + Node.js 24 LTS", "Python 3.14 + Mypy strict"),
    ("Node.js 24 为 LTS", "Python 3.14 为当前后端基线"),
    ("Node.js LTS", "Python 3.14"),
    ("Node.js Releases", "Python Release Cycle"),
    ("OpenTelemetry for Node.js", "OpenTelemetry for Python"),
    ("ClickHouse — JavaScript Client", "ClickHouse — Python Integration"),
    ("TypeScript monorepo", "Python backend package"),
    ("TypeScript 图协议", "Python 图协议"),
    ("TypeScript 任务方案契约", "Python 任务方案契约"),
    ("TypeScript 契约骨架", "Python 契约骨架"),
    ("TypeScript 的现场协议骨架", "Python 的现场协议骨架"),
    ("TypeScript 编译基线", "Python 类型检查基线"),
    ("TypeScript / SQL 契约", "Python / SQL 契约"),
    ("Node.js / TypeScript 总体架构", "Python 3.14 总体架构"),
    ("Node/TS 洞察投影架构", "Python 洞察投影架构"),
    ("Node/TS、投影", "Python、投影"),
    ("NODE & TYPESCRIPT", "PYTHON"),
    ("TYPESCRIPT CONTRACTS", "PYTHON CONTRACTS"),
    ("TYPESCRIPT CONTRACT", "PYTHON CONTRACT"),
    ("TYPESCRIPT", "PYTHON"),
    ("Node Worker", "Python Worker"),
    ("TypeScript 服务", "Python 服务"),
    ("TypeScript", "Python"),
    ("Node.js", "Python 3.14"),
    ("Node/TS", "Python"),
    ("RunTask", "ExecutionUnit"),
    ("execution item", "ExecutionUnit"),
    ("CaseRun / Attempt", "ExecutionUnit / UnitAttempt"),
    ("CaseRun", "ExecutionUnit"),
    ("UnitUnitAttempt", "UnitAttempt"),
    ("attemptId", "unitAttemptId"),
    ("Attempt Seal", "AttemptSeal"),
)

DOCUMENT_TEXT_REPLACEMENTS: dict[str, tuple[tuple[str, str], ...]] = {
    "Atlas_AI用例与浏览器Agent工作流功能设计及实现对齐稿_v0.3.docx": (
        ("功能设计及实现对齐稿  v0.2", "功能设计及实现对齐稿  v0.3"),
        ("v0.2 已对齐  ", "v0.3 已对齐  "),
    ),
    "Atlas_任务中心与批量执行控制面功能设计及实现对齐稿_v0.2.docx": (
        ("功能设计及实现对齐稿  v0.3", "功能设计及实现对齐稿  v0.2"),
    ),
}

REGEX_REPLACEMENTS: tuple[tuple[re.Pattern[str], str], ...] = (
    (re.compile(r"(?<!unit_)attempt_id"), "unit_attempt_id"),
    (
        re.compile(r"(?:retry )*attempts、budget exhausted"),
        "retry attempts、budget exhausted",
    ),
    (re.compile(r"(?<!unit_)attempt\.snapshot"), "unit_attempt.snapshot"),
    (
        re.compile(r"(?<!unit_)attempt/control/page/action/oracle \+ cursor"),
        "unit_attempt/control/page/action/oracle + cursor",
    ),
    (
        re.compile(r"(?<!unit_)attempt:(view|operate|control|comment)"),
        r"unit_attempt:\1",
    ),
    (re.compile(r"(?<!unit-)attempt-scoped"), "unit-attempt-scoped"),
    (re.compile(r"(?<!UNIT_)ATTEMPT_TERMINAL"), "UNIT_ATTEMPT_TERMINAL"),
    (re.compile(r"(?<!unit_)attempt_result_fact"), "unit_attempt_result_fact"),
    (
        re.compile(r"(?<!unit_)attempt\.seal_accepted"),
        "unit_attempt.seal_accepted",
    ),
    (re.compile(r"(?<![A-Za-z])AttemptProjection"), "UnitAttemptProjection"),
    (re.compile(r"(?<![A-Za-z])Attempt Deck"), "UnitAttempt Deck"),
    (re.compile(r"(?<![A-Za-z])AttemptResultFact"), "UnitAttemptResultFact"),
    (re.compile(r"(?<![A-Za-z])Attempt(?![A-Za-z])"), "UnitAttempt"),
)


def iter_paragraphs(document: DocumentType) -> list[Paragraph]:
    """Return body and table paragraphs in document order groups."""
    paragraphs = list(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    return paragraphs


def replace_text(text: str, document_name: str) -> str:
    """Apply stack and terminology replacements to one text value."""
    updated = text
    for source, target in TEXT_REPLACEMENTS:
        updated = updated.replace(source, target)
    for source, target in DOCUMENT_TEXT_REPLACEMENTS.get(document_name, ()):
        updated = updated.replace(source, target)
    for pattern, target in REGEX_REPLACEMENTS:
        updated = pattern.sub(target, updated)
    return updated


def replace_text_nodes(paragraph: Paragraph, document_name: str) -> bool:
    """Replace every OOXML text node, including text inside hyperlinks."""
    changed = False
    for node in paragraph._p.xpath(".//w:t"):
        if node.text is None:
            continue
        updated = replace_text(node.text, document_name)
        if updated != node.text:
            node.text = updated
            changed = True
    return changed


def replace_full_paragraph(paragraph: Paragraph, updated: str) -> None:
    """Replace a full paragraph while preserving its paragraph style."""
    if paragraph.runs:
        paragraph.runs[0].text = updated
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(updated)


def migrate_document(path: Path) -> int:
    """Apply code, terminology, stack, and metadata updates to one document."""
    document = Document(path)
    code_updates = CODE_REPLACEMENTS.get(path.name, {})
    for index, replacement in code_updates.items():
        paragraph = document.paragraphs[index]
        if paragraph.style.name != "Code Block":
            raise RuntimeError(f"Expected Code Block at {path.name}:{index}")
        replace_full_paragraph(paragraph, replacement)

    text_updates = 0
    if path.name in COVER_DATE_DOCUMENTS:
        for row in document.tables[0].rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    changed = False
                    for node in paragraph._p.xpath(".//w:t"):
                        if node.text in {"2026-07-11", "2026-07-12"}:
                            node.text = "2026-07-13"
                            changed = True
                    if changed:
                        text_updates += 1

    for paragraph in iter_paragraphs(document):
        original = paragraph.text
        changed = replace_text_nodes(paragraph, path.name)
        expected = replace_text(original, path.name)
        if paragraph.text != expected:
            # Cover metadata is sometimes split across multiple styled runs.
            replace_full_paragraph(paragraph, expected)
            changed = True
        if changed:
            text_updates += 1

    properties = document.core_properties
    title, status = VERSION_METADATA[path.name]
    properties.title = title
    properties.comments = f"Status: {status}. Backend baseline: Python 3.14."
    properties.last_modified_by = "Atlas Test Space"
    properties.modified = datetime(2026, 7, 13, 12, 0, tzinfo=UTC)
    properties.keywords = "Atlas; Python 3.14; FastAPI; Pydantic; Temporal; Playwright"
    try:
        properties.revision = int(properties.revision or 0) + 1
    except ValueError:
        properties.revision = 1

    document.save(path)
    return text_updates


def main() -> None:
    """Migrate every Atlas DOCX design document in place."""
    for path in sorted(DOCUMENTS_DIR.glob("*.docx")):
        text_updates = migrate_document(path)
        print(
            f"{path.name}: code={len(CODE_REPLACEMENTS.get(path.name, {}))} "
            f"text={text_updates}"
        )


if __name__ == "__main__":
    main()
